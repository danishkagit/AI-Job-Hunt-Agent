import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_JUSTIFY, TA_LEFT

# Colors for professional design
COLOR_PRIMARY_HEX = "#1A365D"  # Deep Navy
COLOR_SECONDARY_HEX = "#4A5568" # Charcoal/Slate
COLOR_TEXT_HEX = "#2D3748"      # Off-black/dark grey

RGB_PRIMARY = RGBColor(0x1A, 0x36, 0x5D)
RGB_SECONDARY = RGBColor(0x4A, 0x55, 0x68)
RGB_TEXT = RGBColor(0x2D, 0x37, 0x48)

def remove_table_borders(table):
    """
    XML helper to completely remove borders from a python-docx table.
    """
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    new_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        new_borders.append(border)
    tblPr.append(new_borders)

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """
    Sets specific cell margins (padding) in twentieths of a point (dxa).
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# --- DOCX CV Generation ---
def generate_cv_docx(cv_data, filepath):
    doc = docx.Document()
    
    # 0.75-inch margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Name & Contact
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(cv_data["name"])
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = RGB_PRIMARY
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run(cv_data["title"])
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(13)
    title_run.font.bold = True
    title_run.font.color.rgb = RGB_SECONDARY
    
    contact = cv_data["contact"]
    contact_parts = [
        contact.get("email"),
        contact.get("phone"),
        contact.get("location"),
        contact.get("linkedin"),
        contact.get("github")
    ]
    contact_str = " | ".join([p for p in contact_parts if p])
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(12)
    contact_run = contact_p.add_run(contact_str)
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = RGB_SECONDARY

    def add_section_heading(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title_text.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGB_PRIMARY
        
        # Add a thin line under the heading using paragraph XML border or keep it clean.
        # Clean heading with color works great for ATS.
        
    # Professional Summary
    add_section_heading("Professional Summary")
    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_after = Pt(6)
    summary_p.paragraph_format.line_spacing = 1.15
    run = summary_p.add_run(cv_data["summary"])
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGB_TEXT
    
    # Skills
    add_section_heading("Technical Skills")
    for category, skill_list in cv_data["skills"].items():
        skills_p = doc.add_paragraph()
        skills_p.paragraph_format.space_after = Pt(4)
        skills_p.paragraph_format.line_spacing = 1.15
        
        cat_run = skills_p.add_run(f"{category}: ")
        cat_run.font.name = 'Calibri'
        cat_run.font.size = Pt(10)
        cat_run.font.bold = True
        cat_run.font.color.rgb = RGB_TEXT
        
        skills_run = skills_p.add_run(", ".join(skill_list))
        skills_run.font.name = 'Calibri'
        skills_run.font.size = Pt(10)
        skills_run.font.color.rgb = RGB_TEXT

    # Professional Experience
    add_section_heading("Professional Experience")
    for job in cv_data["experience"]:
        # 1x2 Table for Title + Date
        table = doc.add_table(rows=1, cols=2)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(5.2)
        table.columns[1].width = Inches(1.8)
        remove_table_borders(table)
        
        # Left cell
        cell_l = table.cell(0, 0)
        set_cell_margins(cell_l, top=60, bottom=60, left=0, right=0)
        p_l = cell_l.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(2)
        run_role = p_l.add_run(job["role"])
        run_role.font.name = 'Calibri'
        run_role.font.size = Pt(11)
        run_role.font.bold = True
        run_role.font.color.rgb = RGB_TEXT
        
        run_sep = p_l.add_run(" – ")
        run_sep.font.name = 'Calibri'
        run_sep.font.size = Pt(11)
        run_sep.font.color.rgb = RGB_SECONDARY
        
        run_co = p_l.add_run(job["company"])
        run_co.font.name = 'Calibri'
        run_co.font.size = Pt(11)
        run_co.font.bold = True
        run_co.font.color.rgb = RGB_SECONDARY
        
        # Right cell
        cell_r = table.cell(0, 1)
        set_cell_margins(cell_r, top=60, bottom=60, left=0, right=0)
        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_r.paragraph_format.space_after = Pt(2)
        run_dur = p_r.add_run(job["duration"])
        run_dur.font.name = 'Calibri'
        run_dur.font.size = Pt(10)
        run_dur.font.bold = True
        run_dur.font.color.rgb = RGB_SECONDARY
        
        # Bullets
        for bullet in job["bullets"]:
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(3)
            bp.paragraph_format.line_spacing = 1.15
            bp_run = bp.add_run(bullet)
            bp_run.font.name = 'Calibri'
            bp_run.font.size = Pt(10)
            bp_run.font.color.rgb = RGB_TEXT
            
    # Selected Projects
    add_section_heading("Selected Projects")
    for proj in cv_data.get("projects", []):
        proj_p = doc.add_paragraph()
        proj_p.paragraph_format.space_before = Pt(4)
        proj_p.paragraph_format.space_after = Pt(2)
        proj_p.paragraph_format.keep_with_next = True
        
        run_name = proj_p.add_run(proj["name"])
        run_name.font.name = 'Calibri'
        run_name.font.size = Pt(10.5)
        run_name.font.bold = True
        run_name.font.color.rgb = RGB_TEXT
        
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.space_after = Pt(4)
        desc_p.paragraph_format.line_spacing = 1.15
        run_desc = desc_p.add_run(proj["description"])
        run_desc.font.name = 'Calibri'
        run_desc.font.size = Pt(10)
        run_desc.font.color.rgb = RGB_TEXT
        
    # Education
    add_section_heading("Education")
    for edu in cv_data["education"]:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(5.2)
        table.columns[1].width = Inches(1.8)
        remove_table_borders(table)
        
        # Left
        cell_l = table.cell(0, 0)
        set_cell_margins(cell_l, top=60, bottom=60, left=0, right=0)
        p_l = cell_l.paragraphs[0]
        p_l.paragraph_format.space_after = Pt(0)
        
        run_deg = p_l.add_run(edu["degree"])
        run_deg.font.name = 'Calibri'
        run_deg.font.size = Pt(10.5)
        run_deg.font.bold = True
        run_deg.font.color.rgb = RGB_TEXT
        
        run_sep = p_l.add_run(" – ")
        run_sep.font.name = 'Calibri'
        run_sep.font.size = Pt(10.5)
        run_sep.font.color.rgb = RGB_SECONDARY
        
        run_inst = p_l.add_run(edu["institution"])
        run_inst.font.name = 'Calibri'
        run_inst.font.size = Pt(10)
        run_inst.font.color.rgb = RGB_SECONDARY
        
        # Right
        cell_r = table.cell(0, 1)
        set_cell_margins(cell_r, top=60, bottom=60, left=0, right=0)
        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_r.paragraph_format.space_after = Pt(0)
        
        run_dur = p_r.add_run(edu["duration"])
        run_dur.font.name = 'Calibri'
        run_dur.font.size = Pt(10)
        run_dur.font.bold = True
        run_dur.font.color.rgb = RGB_SECONDARY
        
    doc.save(filepath)

# --- PDF CV Generation ---
def generate_cv_pdf(cv_data, filepath):
    # Setup document with 0.5-inch margins for neat single-page layout
    doc = SimpleDocTemplate(
        filepath,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    style_name = ParagraphStyle(
        'CVName',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor(COLOR_PRIMARY_HEX),
        alignment=TA_CENTER,
        spaceAfter=2
    )
    
    style_title = ParagraphStyle(
        'CVTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor(COLOR_SECONDARY_HEX),
        alignment=TA_CENTER,
        spaceAfter=4
    )
    
    style_contact = ParagraphStyle(
        'CVContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor(COLOR_SECONDARY_HEX),
        alignment=TA_CENTER,
        spaceAfter=10
    )
    
    style_heading = ParagraphStyle(
        'CVSectionHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor(COLOR_PRIMARY_HEX),
        spaceBefore=10,
        spaceAfter=2,
        keepWithNext=True
    )
    
    style_summary = ParagraphStyle(
        'CVSummary',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=colors.HexColor(COLOR_TEXT_HEX),
        alignment=TA_JUSTIFY,
        spaceAfter=4
    )
    
    style_skills = ParagraphStyle(
        'CVSkills',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor(COLOR_TEXT_HEX),
        spaceAfter=3
    )
    
    style_exp_left = ParagraphStyle(
        'CVExpLeft',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor(COLOR_TEXT_HEX)
    )
    
    style_exp_right = ParagraphStyle(
        'CVExpRight',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor(COLOR_SECONDARY_HEX),
        alignment=TA_RIGHT
    )
    
    style_bullet = ParagraphStyle(
        'CVBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor(COLOR_TEXT_HEX),
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=2
    )
    
    style_proj_title = ParagraphStyle(
        'CVProjTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor(COLOR_TEXT_HEX),
        spaceBefore=3,
        spaceAfter=1,
        keepWithNext=True
    )
    
    style_proj_desc = ParagraphStyle(
        'CVProjDesc',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor(COLOR_TEXT_HEX),
        spaceAfter=4
    )
    
    story = []
    
    # 1. Name & Contact
    story.append(Paragraph(cv_data["name"], style_name))
    story.append(Paragraph(cv_data["title"], style_title))
    
    contact = cv_data["contact"]
    contact_parts = [
        contact.get("email"),
        contact.get("phone"),
        contact.get("location"),
        contact.get("linkedin"),
        contact.get("github")
    ]
    contact_str = " &nbsp;|&nbsp; ".join([p for p in contact_parts if p])
    story.append(Paragraph(contact_str, style_contact))
    
    def add_pdf_section(title_text):
        story.append(Paragraph(title_text.upper(), style_heading))
        story.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor(COLOR_PRIMARY_HEX), spaceBefore=1, spaceAfter=5))

    # 2. Professional Summary
    add_pdf_section("Professional Summary")
    story.append(Paragraph(cv_data["summary"], style_summary))
    
    # 3. Technical Skills
    add_pdf_section("Technical Skills")
    for category, skill_list in cv_data["skills"].items():
        skills_text = f"<b>{category}:</b> {', '.join(skill_list)}"
        story.append(Paragraph(skills_text, style_skills))
        
    # 4. Experience
    add_pdf_section("Professional Experience")
    for job in cv_data["experience"]:
        left_html = f"<b>{job['role']}</b> – <b>{job['company']}</b>"
        right_html = f"{job['duration']}"
        
        # 1x2 Table
        row_data = [
            Paragraph(left_html, style_exp_left),
            Paragraph(right_html, style_exp_right)
        ]
        t = Table([row_data], colWidths=[380, 160])
        t.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING', (0,0), (-1,-1), 1),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(t)
        story.append(Spacer(1, 2))
        
        for bullet in job["bullets"]:
            bullet_text = f"&bull; {bullet}"
            story.append(Paragraph(bullet_text, style_bullet))
        story.append(Spacer(1, 4))
        
    # 5. Selected Projects
    add_pdf_section("Selected Projects")
    for proj in cv_data.get("projects", []):
        story.append(Paragraph(proj["name"], style_proj_title))
        story.append(Paragraph(proj["description"], style_proj_desc))
        
    # 6. Education
    add_pdf_section("Education")
    for edu in cv_data["education"]:
        left_html = f"<b>{edu['degree']}</b> – {edu['institution']}"
        right_html = f"{edu['duration']}"
        row_data = [
            Paragraph(left_html, style_exp_left),
            Paragraph(right_html, style_exp_right)
        ]
        t = Table([row_data], colWidths=[410, 130])
        t.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING', (0,0), (-1,-1), 1),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(t)
        
    doc.build(story)

# --- DOCX Cover Letter Generation ---
def generate_cl_docx(cl_data, filepath):
    doc = docx.Document()
    
    # Standard margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Header
    name_p = doc.add_paragraph()
    name_run = name_p.add_run("Jane Doe")
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.color.rgb = RGB_PRIMARY
    
    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(18)
    contact_run = contact_p.add_run(
        "jane.doe@email.com | +44 7700 900077 | London, UK\n"
        "linkedin.com/in/janedoe | github.com/janedoe"
    )
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = RGB_SECONDARY

    # Date
    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_after = Pt(12)
    date_run = date_p.add_run(cl_data.get("date", "June 4, 2026"))
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(10.5)
    date_run.font.color.rgb = RGB_TEXT
    
    # Recipient
    rec_p = doc.add_paragraph()
    rec_p.paragraph_format.space_after = Pt(14)
    rec_run = rec_p.add_run(
        f"{cl_data.get('recipient', 'Hiring Manager')}\n"
        f"{cl_data.get('company', 'Hiring Team')}"
    )
    rec_run.font.name = 'Calibri'
    rec_run.font.size = Pt(10.5)
    rec_run.font.color.rgb = RGB_TEXT

    # Subject
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    sub_run = sub_p.add_run(cl_data.get("subject", "Application for AI Engineer Position"))
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(11)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGB_PRIMARY
    
    # Salutation
    sal_p = doc.add_paragraph()
    sal_p.paragraph_format.space_after = Pt(10)
    sal_run = sal_p.add_run(cl_data.get("salutation", "Dear Hiring Manager,"))
    sal_run.font.name = 'Calibri'
    sal_run.font.size = Pt(10.5)
    sal_run.font.color.rgb = RGB_TEXT

    # Body Paragraphs
    for para in cl_data.get("paragraphs", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(para)
        r.font.name = 'Calibri'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGB_TEXT
        
    # Sign-off
    so_p = doc.add_paragraph()
    so_p.paragraph_format.space_before = Pt(12)
    so_run = so_p.add_run(cl_data.get("sign_off", "Sincerely,\n\nJane Doe"))
    so_run.font.name = 'Calibri'
    so_run.font.size = Pt(10.5)
    so_run.font.color.rgb = RGB_TEXT
    
    doc.save(filepath)

# --- PDF Cover Letter Generation ---
def generate_cl_pdf(cl_data, filepath):
    doc = SimpleDocTemplate(
        filepath,
        pagesize=letter,
        leftMargin=54,  # 0.75-inch margins for cover letter
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    style_name = ParagraphStyle(
        'CLName',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor(COLOR_PRIMARY_HEX),
        spaceAfter=2
    )
    
    style_contact = ParagraphStyle(
        'CLContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor(COLOR_SECONDARY_HEX),
        spaceAfter=15
    )
    
    style_meta = ParagraphStyle(
        'CLMeta',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor(COLOR_TEXT_HEX),
        spaceAfter=10
    )
    
    style_subject = ParagraphStyle(
        'CLSubject',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor(COLOR_PRIMARY_HEX),
        spaceAfter=12
    )
    
    style_salutation = ParagraphStyle(
        'CLSalutation',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor(COLOR_TEXT_HEX),
        spaceAfter=10
    )
    
    style_body = ParagraphStyle(
        'CLBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14.5,
        textColor=colors.HexColor(COLOR_TEXT_HEX),
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )
    
    style_signoff = ParagraphStyle(
        'CLSignoff',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor(COLOR_TEXT_HEX),
        spaceBefore=12
    )
    
    story = []
    
    # 1. Header
    story.append(Paragraph("Jane Doe", style_name))
    contact_html = "jane.doe@email.com &nbsp;|&nbsp; +44 7700 900077 &nbsp;|&nbsp; London, UK<br/>linkedin.com/in/janedoe &nbsp;|&nbsp; github.com/janedoe"
    story.append(Paragraph(contact_html, style_contact))
    
    # 2. Date
    story.append(Paragraph(cl_data.get("date", "June 4, 2026"), style_meta))
    
    # 3. Recipient
    rec_html = f"{cl_data.get('recipient', 'Hiring Manager')}<br/>{cl_data.get('company', 'Hiring Team')}"
    story.append(Paragraph(rec_html, style_meta))
    
    # 4. Subject
    story.append(Paragraph(cl_data.get("subject", "Application for AI Engineer Position"), style_subject))
    
    # 5. Salutation
    story.append(Paragraph(cl_data.get("salutation", "Dear Hiring Manager,"), style_salutation))
    
    # 6. Body Paragraphs
    for para in cl_data.get("paragraphs", []):
        story.append(Paragraph(para, style_body))
        
    # 7. Sign-off
    so_html = cl_data.get("sign_off", "Sincerely,<br/><br/>Jane Doe").replace("\n", "<br/>")
    story.append(Paragraph(so_html, style_signoff))
    
    doc.build(story)
